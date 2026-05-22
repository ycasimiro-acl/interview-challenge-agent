# -*- coding: utf-8 -*-
"""
generate_cencosud_orquestacion.py
Genera el challenge técnico para Cencosud /
Especialista en Orquestación y Automatización de Procesos Semi Senior.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = (
    r"C:\Users\Yuniesky\Documents\D\D Vit Salva\ACL\interview-challenge-agent"
    r"\Cencosud\Especialista en Orquestación y Automatización Semi Senior"
    r"\challenge-especialista-orquestacion-automatizacion-semisenior.doc"
)

CONTENT = r"""=========================================================================
  CHALLENGE TÉCNICO — ESPECIALISTA EN ORQUESTACIÓN Y AUTOMATIZACIÓN
                      DE PROCESOS SEMI SENIOR
  Cliente   : Cencosud
  Nivel     : Semi Senior
  Stack     : Control-M · Rundeck · CA Workload Automation (CA WA) ·
              UC4 · PagerDuty · JAM ·
              Shell Scripting · PowerShell ·
              Unix/Linux · Windows Server 2008-2025 ·
              Apache Airflow · Apache NiFi (deseable) ·
              ITIL 4 · COBIT (deseable)
  Duración  : 60 minutos
  Modalidad : Mixto (Teórico + Situacional + Código)
  Fecha     : 2026-05-22
  KAM       : Miguel Flores
  Reclutador: Maire Magallanes
  Jefatura  : Jordan Reyes
=========================================================================

⚠  INSTRUCCIÓN PARA EL ENTREVISTADOR — DISEÑO ANTI-IA
════════════════════════════════════════════════════════
Este challenge está diseñado para que NO sea posible responder
correctamente usando un asistente de IA en tiempo real.
Estrategias aplicadas:

1. PREGUNTAS DE EXPERIENCIA REAL
   Cada pregunta técnica exige que el candidato cite un proyecto
   concreto (empresa, contexto, problema real). La IA genera
   respuestas genéricas; el entrevistador presiona con follow-ups
   que verifican autenticidad.

2. FOLLOW-UP SORPRESA VERBAL
   Cada pregunta incluye un FOLLOW-UP que el entrevistador hace
   VERBALMENTE mientras el candidato responde. El candidato no los
   conoce de antemano y no puede prepararlos con IA.

3. DISEÑO EN VIVO CON INTERRUPCIONES
   Las preguntas de diseño de mallas requieren pensamiento en tiempo
   real mientras el entrevistador impone restricciones nuevas a
   medida que el candidato habla.

4. CÓDIGO SIN HERRAMIENTAS
   El code challenge se hace sin acceso a buscadores, IDE ni
   asistentes de IA. Se escribe en papel o editor básico sin
   autocompletado.

5. VOCABULARIO OPERATIVO DE NICHO
   Control-M, CA WA, UC4 y la semántica de "condiciones" son
   suficientemente específicas para que una respuesta genérica
   de IA sea inmediatamente identificable. Se evalúa vocabulario
   real: ABEND, IN/OUT Conditions, Order Date (Odate), max wait,
   resource pool — términos que no aparecen en tutoriales básicos.

─────────────────────────────────────────────────────────────────────
  MAPA DE COBERTURA — REQUISITOS EXCLUYENTES
─────────────────────────────────────────────────────────────────────
  RE-1  Orquestación empresarial .... Preguntas A-1, A-2, A-3
  RE-2  Scripting Shell/PowerShell .. Pregunta D-1 + Code Challenge
  RE-3  Gestión de servicios ITIL ... Preguntas B-1, B-2
  RE-4  Automatización / diseño ..... Preguntas C-1, C-2
  RE-D  Diferenciadores (deseable) .. Pregunta E-1 (Airflow/NiFi)

═════════════════════════════════════════════════════════════════════
  SECCIÓN 0 — WARM-UP (5 min) · No evaluado con puntaje
═════════════════════════════════════════════════════════════════════

WU-1: "Cuéntame de la malla de procesos más compleja que hayas
       diseñado o mantenido tú personalmente. ¿Qué herramienta
       usabas, cuántos jobs tenía aproximadamente y cuál era el
       proceso de negocio que automatizaba?"

       Señal positiva : nombra empresa, herramienta concreta,
                        número aproximado de jobs, proceso real
                        (cierre contable, carga DWH, extracción SAP).
       Señal negativa : respuesta genérica sin detalles propios,
                        no puede dar un contexto de negocio concreto.

WU-2: "Fuera del horario administrativo, ¿alguna vez atendiste un
       incidente crítico de un job scheduler en producción? ¿Qué
       fue lo primero que hiciste al recibir la alerta?"

       Señal positiva : protocolo real de diagnóstico — revisar log,
                        verificar infra, evaluar impacto antes de actuar.
       Señal negativa : "reinicié el job" como primera acción.

═════════════════════════════════════════════════════════════════════
  SECCIÓN A — FUNDAMENTOS DE ORQUESTACIÓN (15 min) [RE-1]
═════════════════════════════════════════════════════════════════════

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta A-1: Dependencias, condiciones y estados en Control-M
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  En Control-M, ¿cómo defines las dependencias entre jobs dentro
  de una malla? Describe el mecanismo exacto: ¿qué sucede si un
  job upstream falla a la mitad del proceso? ¿Cómo evitas que los
  jobs downstream queden bloqueados indefinidamente?

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - IN/OUT Conditions: el job downstream espera una IN Condition que
    el upstream agrega como OUT Condition al completar exitosamente
  - Si el upstream falla (ABEND o FAILED), NO genera la OUT Condition
    → el downstream queda en estado WAIT indefinidamente
  - Configurar "max wait" en jobs downstream para que entren en
    FAILED si la condición no llega dentro del tiempo esperado
  - Diferencia entre FAILED (terminó con error esperado) y ABEND
    (terminación anómala/abrupta) — impacto diferente en el flujo

  EXPLICACIÓN DE REFERENCIA (buenas prácticas):
  En Control-M las dependencias se modelan mediante condiciones de
  prerequisito (IN Conditions) y de salida (OUT Conditions). El pool
  de condiciones activas es global por Odate (Order Date). Si el
  upstream no completa exitosamente, su OUT Condition no se agrega
  al pool y el downstream queda en WAIT. Un riesgo adicional es la
  contaminación entre Odates: si una condición quedó activa del día
  anterior (por ejecución manual o Force Order), el downstream del
  día siguiente puede dispararse incorrectamente. La práctica
  correcta combina: (1) max wait para que el downstream no espere
  eternamente; (2) alerta en PagerDuty para estados WAIT que
  excedan el SLA; (3) job de "limpieza de condiciones" al inicio
  de cada Odate para eliminar condiciones huérfanas del día anterior.

  RESPUESTA DESTACADA (diferenciador senior):
  Menciona Cyclic jobs y cómo sus condiciones se comportan diferente
  de los jobs regulares. Describe el riesgo de condiciones
  "permanentes" (sin fecha de expiración) vs condiciones vinculadas
  al Odate. Explica cómo hacer Force-ordering manual de un job con
  sus condiciones para resolver dependencias colgadas sin correr
  todo el flujo desde cero.

  RED FLAGS:
  - No distingue entre FAILED y ABEND (no ha trabajado con Control-M
    real; son estados con semánticas y consecuencias distintas).
  - "Los downstream se caen automáticamente" — incorrecto en Control-M.
  - Confunde dependencias de job (condiciones) con reglas de
    scheduling (calendarios/crons).
  - No conoce el concepto de Odate (Order Date).

  FOLLOW-UP SORPRESA (preguntar verbalmente mientras responde):
  "¿Qué pasa con las condiciones si haces un Force Order manual
   de un job fuera del flujo habitual? ¿Cómo evitas que eso
   contamine la ejecución del día siguiente de otros jobs que
   esperaban esa condición?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No explica dependencias o da respuesta genérica      |
  | 1 - Básico    | Sabe que hay dependencias pero no explica condiciones |
  | 2 - Competente| IN/OUT Conditions + WAIT + max wait + ABEND vs FAILED|
  | 3 - Destacado | + Odate + contaminación entre días + Force Order      |
  +----------------+------------------------------------------------------+

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta A-2: Distribución de carga y resource pools
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Cencosud opera a nivel regional con múltiples países y cierres
  contables distintos. Tienes 400 jobs nocturnos que deben terminar
  antes de las 06:00 AM hora Chile. Algunos jobs compiten por el
  mismo servidor FTP y la misma base de datos SAP que soporta
  máximo 10 conexiones simultáneas. ¿Cómo distribuyes la carga?
  ¿Qué parámetros configuras en tu herramienta de orquestación para
  evitar cuellos de botella sin alargar la ventana de proceso?
  [El entrevistador puede interrumpir mientras el candidato habla.]

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - Time constraints: Not Earlier Than / Not Later Than / Must End By
  - Resource pools (semáforos): limitar concurrencia al servidor FTP
    y a SAP — ej: resource pool "SAP_CONN" con capacidad 10
  - Análisis del critical path para identificar la cadena de
    dependencias más larga que determina la duración total
  - Staggering: distribuir inicio de jobs no críticos para evitar
    picos de carga simultáneos al arranque
  - Agent groups por zona geográfica para ejecutar en el agente
    del país correspondiente sin cambiar la lógica de la malla

  EXPLICACIÓN DE REFERENCIA:
  El dimensionamiento empieza por el critical path: identificar la
  cadena de dependencias más larga que determina el tiempo mínimo
  total. Control-M y CA WA permiten definir Resource Pools que
  actúan como semáforos: el pool "SAP_CONN" con capacidad 10
  garantiza que máximo 10 jobs intenten conectar a SAP en paralelo;
  los demás esperan en cola sin fallar. Las ventanas se modelan
  con time constraints: Must End By activa una alerta o fallo si
  el job no puede terminar antes de la hora límite; Not Later Than
  es el umbral de inicio máximo para que el scheduler pueda
  garantizar que el job termine a tiempo. Los agent groups permiten
  una operación regional transparente para la malla.

  RESPUESTA DESTACADA (diferenciador senior):
  Menciona la funcionalidad de forecasting/simulation de Control-M
  para estimar la duración total antes de comprometerse con un SLA
  nuevo. Habla de KPIs de duración promedio por job para detectar
  degradación antes de que rompa la ventana. Propone una estrategia
  de "shakeout" en staging con datos representativos antes de
  modificar ventanas en producción.

  RED FLAGS:
  - "Ponemos más servidores" sin hablar de la herramienta de
    orquestación ni de concurrencia controlada.
  - No conoce resource pools o semáforos de concurrencia.
  - No menciona time constraints (habla solo de "crons").
  - No puede explicar qué es el critical path de una malla.

  FOLLOW-UP SORPRESA:
  "El job de cierre contable de Perú siempre tardaba 25 minutos.
   Esta semana tardó 80 minutos y casi rompió la ventana. Terminó
   exitoso, sin ningún error. ¿Qué investigas y cómo lo detectas
   a tiempo la próxima vez antes de que sea tarde?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | Sin experiencia con cargas masivas o respuesta vaga  |
  | 1 - Básico    | Sabe que hay límites de concurrencia pero no configura|
  | 2 - Competente| Resource pools + time constraints + critical path    |
  | 3 - Destacado | + forecasting + KPIs de duración + shakeout staging  |
  +----------------+------------------------------------------------------+

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta A-3: Diferencias operativas entre herramientas (no marketing)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Has trabajado con más de una herramienta de orquestación
  (Control-M, Rundeck, CA WA/AutoSys). Dame 3 diferencias concretas
  y operativas que experimentaste tú mismo — no del marketing de los
  vendors, sino de lo que te complicó o facilitó el trabajo real.
  ¿Para qué empresa o caso de uso elegirías cada una?

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - Control-M tiene calendarios de negocio nativos (días hábiles,
    festivos por país) — Rundeck no tiene eso por defecto
  - CA WA (AutoSys) define jobs en JIL (Job Information Language),
    texto plano que facilita versionado en Git y gestión masiva por
    scripting; Control-M usa GUI + CLI (Helix Control-M API)
  - Rundeck está orientado a runbooks de operaciones (infraestructura);
    Control-M y CA WA son más maduros para batch empresarial masivo
  - Modelo de agentes diferente: Control-M Agent instala componentes
    propietarios en cada servidor; Rundeck Node es más liviano

  EXPLICACIÓN DE REFERENCIA:
  Control-M (BMC) es líder en batch scheduling enterprise con
  calendarios de negocio nativos, GUI de monitoring masivo y
  conceptos propios (Conditions Pool, Odate, Cyclic). CA WA
  (antes AutoSys) tiene herencia fuerte en mainframe/UNIX y usa JIL
  para definir jobs en texto plano — la gestión masiva por scripting
  es una ventaja real para catálogos grandes. Rundeck es ideal para
  automatización de runbooks operativos (no batch financiero): sus
  jobs son más parecidos a scripts Ansible que a procesos de cierre
  contable. La elección depende del volumen de batch, complejidad
  de calendarios de negocio, integración con el ecosistema existente
  y el presupuesto de licenciamiento.

  RESPUESTA DESTACADA:
  Menciona experiencia migrando entre herramientas y los desafíos
  de mapear conceptos (las IN/OUT Conditions de Control-M no tienen
  equivalente directo en CA WA). Habla del modelo de licenciamiento
  por agente/tarea vs open-source y cómo afecta las decisiones de
  escala. Comenta sobre la API REST de Helix Control-M para integrar
  el scheduler con pipelines CI/CD modernos.

  RED FLAGS:
  - Comparativa genérica extraída de artículos sin experiencia propia.
  - Solo conoce una herramienta y no puede comparar.
  - Confunde Rundeck con Jenkins o con Ansible Tower.

  FOLLOW-UP SORPRESA:
  "Si hoy tuvieras que migrar 500 jobs de CA WA a Control-M, ¿cuál
   sería el primer problema técnico que enfrentarías? ¿Qué concepto
   de CA WA no tiene equivalente directo en Control-M?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | Solo conoce una herramienta o da resp. de marketing  |
  | 1 - Básico    | Conoce ambas superficialmente                        |
  | 2 - Competente| 3 diferencias operativas + criterio de selección real |
  | 3 - Destacado | + experiencia migración + API moderna + licenciamiento|
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN B — INCIDENT RESPONSE E ITIL (10 min) [RE-3]
═════════════════════════════════════════════════════════════════════

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta B-1: Incident response — job crítico caído a las 03:15 AM
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Son las 03:15 AM. Recibes una alerta de PagerDuty: el job de
  cierre contable de Chile falló en el paso 7 de 12. Los 6 pasos
  anteriores ya ejecutaron y algunos tocaron la base de datos.
  No existe un runbook actualizado para este job específico.
  El jefe de operaciones te llama en 10 minutos.
  Descríbeme con exactitud qué haces en esos 10 minutos: qué
  herramientas abres, qué información recopilas y qué decisiones
  tomas antes de esa llamada.

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - Revisar output/log del job en el scheduler ANTES de cualquier
    acción: código de retorno (exit code/return code), stderr,
    últimas líneas del log del script
  - Identificar la categoría del fallo: ¿infra (servidor caído,
    tablespace lleno, timeout de red) o lógica (error en el script)?
  - Evaluar idempotencia: ¿los pasos anteriores pueden correr dos
    veces sin duplicar datos? No reiniciar sin esta información
  - Brief de 3 líneas para la llamada: qué falló, dónde, qué
    opciones hay (retry / rollback / escalar a aplicaciones)
  - Documentar el timeline desde la alerta para el post-mortem

  EXPLICACIÓN DE REFERENCIA:
  El principio fundamental: diagnóstico primero, acción después.
  Los 10 minutos se distribuyen: (1) abrir el viewer de Control-M
  Monitoring — exit code, stderr, últimas líneas del log; un exit
  127 indica "comando no encontrado" (problema de PATH), exit 1
  con mensaje de BD indica error transaccional, exit -1 es timeout
  o kill; (2) verificar si el servidor del agente responde (SSH/RDP);
  (3) revisar si hay transacciones pendientes en DB — sin esta
  información, un retry puede duplicar datos financieros o corromper
  el cierre; (4) preparar el brief: "Falló el paso 7 con exit code 1
  por timeout de conexión a DB. Opciones: (a) retry si DBA confirma
  ausencia de transacciones abiertas, (b) escalar a aplicaciones si
  el error es de lógica, (c) ejecutar pasos 7-12 manualmente si se
  confirma idempotencia."

  RESPUESTA DESTACADA:
  Menciona documentar el timeline completo del incidente (hora alerta,
  diagnóstico, acción, resolución) para el post-mortem. Habla del
  riesgo de scripts que usan "set -e" vs sin él en relación al exit
  code reportado al scheduler — un script puede swallowear el error
  interno y reportar exit 0. Propone contactar al DBA de guardia
  antes de cualquier retry sobre datos financieros. Diferencia
  entre MTTR (Mean Time to Restore el servicio) y la causa raíz.

  RED FLAGS:
  - "Primero intento reiniciar el job" sin diagnóstico — señal
    crítica de falta de criterio con datos transaccionales.
  - No menciona revisar logs antes de llamar al jefe.
  - No distingue entre fallo de infraestructura y fallo de lógica.
  - No evalúa idempotencia antes de proponer un retry.

  FOLLOW-UP SORPRESA (interrumpir mientras responde):
  "El log muestra: 'ORA-01555: snapshot too old'. No es un error
   de red ni de infraestructura. ¿Qué significa ese error Oracle?
   ¿Cómo cambia tu decisión de hacer retry?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | Primera acción es retry sin diagnóstico              |
  | 1 - Básico    | Revisa logs pero no evalúa idempotencia ni impacto   |
  | 2 - Competente| Diagnóstico + idempotencia + brief estructurado      |
  | 3 - Destacado | + timeline post-mortem + DBA antes del retry +ORA err|
  +----------------+------------------------------------------------------+

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta B-2: Change Management — migración de plataforma
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Debes migrar 200 jobs de Control-M versión 9 a Control-M versión 21
  en el ambiente productivo de Cencosud. El cambio ocurre en fin de
  semana. El mes anterior hubo un incidente de 4 horas causado por
  otro cambio de plataforma. ¿Cómo estructuras el RFC (Request for
  Change)? ¿Cuál es tu plan de rollback? ¿Cómo generas confianza
  con el área de negocio antes del cambio?

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - RFC con: descripción, justificación, impacto, plan de
    implementación paso a paso, plan de rollback, ventana de
    mantenimiento y criterios de éxito/fallo explícitos
  - Pruebas en ambiente no-productivo (staging/QA) con datos reales
    antes del cambio productivo
  - Rollback plan específico: criterio de activación con umbral
    concreto (ej: "si el 20% de jobs críticos no enciende en 30 min
    → rollback") y procedimiento técnico detallado
  - Comunicación proactiva al negocio: cronograma, riesgos residuales,
    contact person durante la ventana de mantenimiento
  - Aprobación del CAB para un cambio Mayor de esta envergadura

  EXPLICACIÓN DE REFERENCIA:
  Una migración de versión de scheduler es un cambio Mayor (Major
  Change) en ITIL, requiriendo aprobación del CAB. El plan incluye:
  (1) inventario con criticidad de negocio por job — los críticos
  se prueban primero en staging; (2) checklist de validación post-
  migración por tipo de job (scripts shell, PowerShell, FTP, SAP);
  (3) criterio de éxito claro: "100% de jobs críticos encendiendo
  correctamente 45 min después del inicio"; (4) criterio de rollback:
  "si un job crítico no enciende pasados 45 min → rollback inmediato";
  (5) ensayo del procedimiento de rollback en staging antes del cambio
  real. La confianza se construye con evidencia: resultados del testing
  en staging + lista de riesgos residuales + transparencia sobre qué
  se hace si algo falla durante la ventana.

  RESPUESTA DESTACADA:
  Propone migración por fases (no big bang): primero los jobs de menor
  criticidad para validar el proceso, luego los críticos. Menciona el
  concepto de "parallel run" — ejecutar ambas versiones en paralelo
  durante un período de validación. Habla de incompatibilidades de
  definición de jobs entre versiones de Control-M (algunos parámetros
  y el modelo de agents cambian entre major versions). Propone usar
  el post-mortem del incidente anterior para incorporar lecciones
  aprendidas en este cambio.

  RED FLAGS:
  - No tiene un plan de rollback concreto — señal crítica.
  - "Hago los cambios y aviso después" — falta de proceso ITIL.
  - El rollback es solo "volver a la versión anterior" sin criterio
    de activación ni procedimiento técnico.
  - No considera el contexto del incidente anterior para manejar
    expectativas del negocio.

  FOLLOW-UP SORPRESA:
  "A las 02:00 AM durante la migración descubres que 15 jobs que
   usaban el parámetro 'cyclic_type' de Control-M 9 no tienen
   equivalente directo en la versión 21. ¿Cómo lo manejas sin
   detener el cambio completo ni dejar esos jobs sin ejecutar?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No tiene proceso estructurado para cambio mayor      |
  | 1 - Básico    | Menciona RFC pero sin rollback plan ni criterios      |
  | 2 - Competente| RFC + rollback con criterios + staging + comunicación |
  | 3 - Destacado | + migración por fases + parallel run + post-mortem   |
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN C — DISEÑO DE MALLAS EN VIVO (15 min) [RE-4]
═════════════════════════════════════════════════════════════════════

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta C-1: Diseño en vivo con interrupciones
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  El área de Finanzas de Cencosud te pide automatizar este proceso:
  cada día hábil a las 22:00 deben ejecutarse 3 extracciones desde
  SAP (8-25 min cada una), luego una consolidación que necesita las
  3 fuentes, luego 5 reportes independientes que usan el dato
  consolidado, y finalmente un envío de email con los reportes.
  La ventana disponible es hasta las 01:00 AM.

  Diseña la malla describiendo: qué ejecuta en paralelo, qué en
  secuencia, cómo manejas el fallo de una extracción SAP y qué
  pasa si el email falla.
  [Interrumpir con restricciones nuevas mientras el candidato habla.]

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - Las 3 extracciones SAP van en PARALELO (fan-out al inicio)
  - La consolidación depende de las 3 extracciones (fan-in) —
    no ejecuta si falta alguna fuente
  - Los 5 reportes van en PARALELO tras la consolidación (fan-out)
  - El email es el paso final, depende de los 5 reportes
  - Fallo extracción SAP: NO ejecutar consolidación, alertar al
    negocio, no contaminar la ejecución del día siguiente
  - Fallo email: es idempotente → retry automático (ej: 3 intentos,
    5 min entre intentos) antes de escalar la alerta

  EXPLICACIÓN DE REFERENCIA:
  Patrón fan-out → fan-in → fan-out clásico en batch orchestration.
  Las 3 extracciones en paralelo reducen el tiempo total: en secuencia
  serían hasta 8+25+20=53 min; en paralelo son 25 min máximo. La
  consolidación tiene una IN Condition que depende de las 3 OUT
  Conditions. Si falta alguna, no ejecuta. Los 5 reportes tienen
  cada uno una IN Condition de la consolidación. Para el fallo de
  extracción: en un proceso financiero no se asume dato parcial como
  válido — la consolidación no corre sin las 3 fuentes. Para el
  email: es idempotente (enviarlo dos veces no corrompe datos), por
  lo que se configura retry en el scheduler. Incluir antes del email
  una validación de tamaño/recuento de archivos como control de
  calidad del dato.

  RESPUESTA DESTACADA:
  Agrega un job de validación de completitud después de cada extracción
  (verifica recuento de registros vs. día anterior, alerta si hay
  discrepancia > umbral). Propone retries para las extracciones SAP
  (no solo para el email) con máximo 2 intentos antes de declarar
  fallo. Modela los días hábiles para que el proceso no corra en fin
  de semana ni festivos por país. Verifica que el critical path
  (extracción 25 min + consolidación + reportes en paralelo) cabe
  dentro de la ventana de 3 horas.

  RED FLAGS:
  - Pone las 3 extracciones SAP en SECUENCIA — ineficiente y demuestra
    falta de experiencia con diseño de mallas.
  - No distingue entre jobs idempotentes (email) y no idempotentes
    (extracción + consolidación) al definir la estrategia de retry.
  - Solo describe el happy path; no maneja ningún escenario de fallo.
  - El diseño propuesto no cabe dentro de la ventana de 3 horas.

  INTERRUPCIONES SORPRESA (agregar mientras el candidato diseña):

  Interrupción 1: "El servidor SAP solo soporta 2 conexiones
  simultáneas. ¿Cómo ajustas el diseño?"
  Respuesta esperada: resource pool "SAP_CONN" con capacidad 2;
  2 extracciones corren en paralelo, la 3ra espera en cola.

  Interrupción 2: "El negocio pide que si la extracción de Argentina
  falla, los reportes de Chile y Perú igualmente se generen con la
  data disponible. ¿Cómo cambias el diseño?"
  Respuesta esperada: desacoplar la IN Condition de la consolidación
  para aceptar completitud parcial — esto cambia el SLA del proceso
  y requiere validación del negocio sobre la calidad del dato.

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No puede diseñar la malla o pone todo en secuencia   |
  | 1 - Básico    | Identifica el paralelismo pero no maneja fallos       |
  | 2 - Competente| Fan-in/fan-out + manejo diferenciado de fallos        |
  | 3 - Destacado | + validación de dato + retries SAP + ajuste a restric.|
  +----------------+------------------------------------------------------+

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta C-2: Observabilidad y KPIs de la plataforma
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Con 400 jobs nocturnos en producción, ¿cómo defines y monitoreas
  la "salud" de la plataforma de orquestación? Dame los 3 KPIs más
  importantes que reportarías al área de negocio, y describe cómo
  los calculas y desde dónde obtienes los datos.

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - SLA Compliance: % de jobs que terminan antes del Must End By
    → datos: histórico de end_time vs Must End By por Odate
  - Job Failure Rate: % de jobs que entran en ABEND o FAILED sin
    retorno exitoso → datos: estados finales del scheduler
  - MTTR (Mean Time to Repair): tiempo promedio entre alerta y
    resolución → datos: registros ITSM + timestamps de PagerDuty

  EXPLICACIÓN DE REFERENCIA:
  Los KPIs deben estar vinculados al impacto de negocio, no solo
  a métricas técnicas. SLA Compliance responde: "¿cuántos procesos
  de negocio terminaron a tiempo?" Job Failure Rate responde:
  "¿cuántos procesos requirieron intervención manual?" MTTR
  responde: "cuando algo falla, ¿cuánto tarda el equipo en
  restaurar?" Los datos se extraen del repositorio del scheduler
  (Control-M tiene APIs y BD interna consultable), del ITSM
  (ServiceNow, Jira SM) y de PagerDuty (timestamps de alerta y
  acknowledge). Para presentación gerencial: dashboard mensual con
  tendencia de los 3 KPIs + top 5 jobs con mayor impacto en SLA.

  RESPUESTA DESTACADA:
  Agrega un 4to KPI: Agent Infrastructure Health (disponibilidad
  de agentes y servidores vinculados). Propone alertas proactivas:
  si un job supera el doble de su duración histórica promedio →
  alerta de "posible colgado" antes de llegar al Must End By.
  Diferencia entre monitoreo reactivo (alertar cuando falla) y
  monitoreo proactivo (detectar tendencias de degradación antes
  de fallar). Menciona integración de PagerDuty con el scheduler
  via webhook para alertas automáticas.

  RED FLAGS:
  - Solo menciona "si falla, PagerDuty avisa" — sin KPIs ni
    métricas estructuradas.
  - Los KPIs son puramente técnicos sin vinculación al negocio.
  - No sabe de dónde obtiene los datos para calcular los KPIs.

  FOLLOW-UP SORPRESA:
  "El área de negocio te pide un semáforo (rojo/amarillo/verde)
   del estado de la plataforma que se refresque cada hora. ¿Qué
   criterios defines para cada color y qué herramienta usarías
   para implementarlo sin desarrollar desde cero?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No tiene KPIs o solo menciona alertas reactivas      |
  | 1 - Básico    | Da KPIs técnicos sin vinculación al negocio          |
  | 2 - Competente| 3 KPIs vinculados al negocio + fuente de datos       |
  | 3 - Destacado | + monitoreo proactivo + webhook PagerDuty + dashboard |
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN D — SCRIPTING (10 min) [RE-2]
═════════════════════════════════════════════════════════════════════

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
  Pregunta D-1: Shell Script robusto para producción con Control-M
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ENUNCIADO:
  Escribe (o describe con pseudocódigo detallado) un shell script Bash
  que: (1) recibe el nombre de un archivo de log como parámetro,
  (2) verifica que el archivo existe y no está vacío,
  (3) busca la palabra "ERROR" (case-insensitive) y cuenta ocurrencias,
  (4) si hay más de 5 errores: escribe en un archivo de reporte la
  fecha/hora, el nombre del log y el conteo, y sale con exit code 1,
  (5) si hay 5 o menos errores: sale con exit code 0.
  El script será invocado por Control-M. ¿Qué consideraciones
  adicionales agrega para ese contexto específico?

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - set -euo pipefail al inicio — evitar errores silenciosos
  - Validar que se pasó el parámetro: [ -z "$1" ] && exit 1
  - Verificar que el archivo existe y no está vacío: -f y -s
  - grep -ci "error" para count case-insensitive
  - Exit codes explícitos: 0 para success, no-0 para ABEND —
    Control-M usa el exit code para determinar SUCCESS o ABEND
  - Rutas absolutas para comandos (/usr/bin/grep, /bin/date) porque
    Control-M ejecuta en entorno mínimo sin el PATH del usuario
  - Output de errores a STDERR (>&2), no a STDOUT

  EXPLICACIÓN DE REFERENCIA:
  Los scripts invocados por schedulers tienen una regla crítica:
  el exit code es el lenguaje del scheduler. Control-M interpreta
  exit 0 como SUCCESS, cualquier otro valor como ABEND. Un error
  frecuente: grep sale con exit code 1 cuando NO encuentra el
  patrón (no es un error, es "no encontrado") — hay que diferenciarlo.
  El "set -euo pipefail" hace que cualquier error en el pipeline
  falle el script, evitando falsas ejecuciones exitosas. En Control-M,
  el agente ejecuta scripts en un entorno mínimo sin el .bashrc del
  usuario — las rutas absolutas son críticas. Un lockfile
  ($0.lock) evita ejecuciones concurrentes del mismo script si un
  job tarda más de lo esperado.

  RESPUESTA DESTACADA:
  Menciona el problema del exit code de grep (1 = not found, 2 =
  error real) y cómo diferenciarlo con COUNT=$(grep -ci "error"
  "$FILE" || true). Propone logging estructurado con timestamp para
  trazabilidad en el viewer de Control-M. Habla de lockfile para
  evitar ejecuciones concurrentes. Menciona que el umask puede
  afectar los permisos del archivo de reporte generado.

  RED FLAGS:
  - No usa exit codes explícitos — evidencia de no haber integrado
    scripts con schedulers empresariales.
  - No valida el parámetro de entrada — el job fallará en Control-M
    si no se pasa el argumento.
  - Usa echo "error" y exit 0 indiscriminadamente.
  - No menciona el problema de PATH en el entorno del agente.

  FOLLOW-UP SORPRESA:
  "¿Qué pasa si el archivo de log sigue escribiéndose mientras tu
   script lo lee con grep? ¿Hay riesgo de recuento incorrecto o
   de que el script falle? ¿Cómo lo manejas en producción?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No puede escribir el script o ignora exit codes      |
  | 1 - Básico    | Script funcional pero sin manejo de errores robusto  |
  | 2 - Competente| set -euo + validaciones + exit codes + rutas absolutas|
  | 3 - Destacado | + grep exit code handling + lockfile + logging estruc.|
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN E — DIFERENCIADOR / ECOSISTEMA MODERNO (5 min) [RE-D]
═════════════════════════════════════════════════════════════════════

  Pregunta E-1: Airflow vs Control-M — coexistencia y estrategia
  ─────────────────────────────────────────────────────────────────

  ENUNCIADO:
  Cencosud está evaluando si usar Apache Airflow para flujos de datos
  modernos (pipelines hacia un Data Lake) mientras mantiene Control-M
  para los procesos batch transaccionales críticos. ¿Cuál es tu
  opinión técnica? ¿Para qué caso de uso usarías cada uno?
  ¿Cuáles son los riesgos de tener dos plataformas de orquestación
  coexistiendo?

  RESPUESTA MÍNIMA ESPERADA (must-mention):
  - Airflow es ideal para DAGs de data engineering (ETL/ELT, ML
    pipelines) con dependencias complejas definidas en Python
  - Control-M es superior para batch empresarial con calendarios de
    negocio, grandes volúmenes de jobs, compliance ITIL y SLAs
    financieros críticos
  - Riesgo de coexistencia: fragmentación del monitoreo (2 consolas),
    duplicación de procesos, curva de aprendizaje, dificultad para
    tener una vista unificada de la salud de la plataforma

  EXPLICACIÓN DE REFERENCIA:
  Airflow y Control-M resuelven problemas diferentes. Airflow define
  workflows como código Python (DAGs) — excelente para data pipelines
  con integración nativa al ecosistema big data (Spark, dbt, GCS, S3)
  y un modelo de backfill muy potente. Control-M está diseñado para
  batch transaccional crítico con calendarios de negocio nativos,
  auditoría ITIL y SLAs financieros. La estrategia correcta es
  mantener Control-M para los procesos críticos y adoptar Airflow
  para los pipelines analíticos modernos, con integración vía API
  cuando sea necesario. El riesgo mayor de coexistencia es la
  fragmentación del monitoreo: el equipo debe monitorear dos consolas,
  lo que aumenta el riesgo de que un fallo pase desapercibido.

  RESPUESTA DESTACADA:
  Menciona que Control-M tiene un plugin nativo para Airflow que
  permite orquestar DAGs desde Control-M (vista unificada). Habla
  de Prefect o Dagster como alternativas más modernas a Airflow con
  mejor UX y retries nativos. Propone una arquitectura donde
  Control-M actúa como orquestador maestro que dispara DAGs de
  Airflow como jobs downstream, manteniendo la vista de monitoreo
  centralizada.

  RED FLAGS:
  - "Airflow reemplaza a Control-M completamente" — demuestra no
    entender los casos de uso de cada herramienta.
  - No puede describir qué es un DAG de Airflow.
  - No ve ningún riesgo en tener dos plataformas de orquestación.

  FOLLOW-UP SORPRESA:
  "Si un DAG de Airflow que alimenta datos al proceso de cierre
   contable en Control-M falla a las 21:50, ¿cómo garantizas que
   Control-M detecte ese fallo antes de intentar ejecutar el
   cierre a las 22:00? ¿Cómo conectas ambos mundos?"

  PUNTUACIÓN:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No conoce Airflow o cree que reemplaza a Control-M  |
  | 1 - Básico    | Conoce Airflow pero no puede comparar con fundamento |
  | 2 - Competente| Diferenciación clara de casos de uso + riesgos coex. |
  | 3 - Destacado | + integración Control-M/Airflow + alternativas modern.|
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN F — CODE CHALLENGE POWERSHELL (15 min) [RE-2]
═════════════════════════════════════════════════════════════════════

  INSTRUCCION AL ENTREVISTADOR:
  Entregar el enunciado al candidato (impreso o en pantalla).
  3 minutos para leer y preguntar dudas; 12 minutos para implementar.
  NO puede usar buscadores, IDE ni asistentes de IA.
  Puede escribir en papel o editor básico sin autocompletado.

─────────────────────────────────────────────────────────────────────
  ENUNCIADO (entregar al candidato — no el resto de esta sección):
─────────────────────────────────────────────────────────────────────

  Recibes un archivo "jobs_status.csv" con este formato:

  job_name,start_time,end_time,status,expected_duration_min
  CIERRE_CL_CONTABLE,22:00,22:45,COMPLETED,40
  EXTRACCION_SAP1,22:00,23:10,COMPLETED,30
  CONSOLIDACION_DWH,23:15,,RUNNING,20
  ENVIO_REPORTES,,,WAITING,10
  BACKUP_DB_PROD,21:30,22:10,COMPLETED,25

  Escribe un script PowerShell que:
  1. Lee el CSV desde la ruta en $args[0]; si no existe: exit 1
  2. Para jobs COMPLETED: si la duración real supero el expected,
     imprime: [SLA_BREACH]    <job_name> - Real: X min | SLA: Y min
  3. Para jobs RUNNING: si el tiempo transcurrido desde start_time
     supera el DOBLE del expected_duration_min, imprime:
     [POSSIBLE_HANG] <job_name> - Transcurrido: X min | Límite: Y min
  4. Al final imprime un resumen:
     Jobs procesados: N | Alertas generadas: M
  5. Sale con exit code 1 si hay alertas, 0 si no hay ninguna

─────────────────────────────────────────────────────────────────────
  SOLUCIÓN DE REFERENCIA (solo entrevistador — NO mostrar):
─────────────────────────────────────────────────────────────────────

  $CsvPath = $args[0]
  if (-not (Test-Path $CsvPath)) {
      Write-Error "Archivo no encontrado: $CsvPath"
      exit 1
  }
  $jobs    = Import-Csv $CsvPath
  $now     = Get-Date
  $alertas = 0

  foreach ($job in $jobs) {
      $expected = [int]$job.expected_duration_min

      # Jobs completados: verificar SLA
      if ($job.status -eq "COMPLETED" -and
          $job.start_time -ne "" -and $job.end_time -ne "") {
          $start  = [DateTime]::ParseExact($job.start_time,"HH:mm",$null)
          $end    = [DateTime]::ParseExact($job.end_time,  "HH:mm",$null)
          if ($end -lt $start) { $end = $end.AddDays(1) }  # cruza medianoche
          $actual = [int]($end - $start).TotalMinutes
          if ($actual -gt $expected) {
              Write-Host ("[SLA_BREACH]    {0,-22} - Real: {1} min | SLA: {2} min" `
                  -f $job.job_name, $actual, $expected)
              $alertas++
          }
      }

      # Jobs en ejecución: verificar posible colgado
      if ($job.status -eq "RUNNING" -and $job.start_time -ne "") {
          $startDt = [DateTime]::ParseExact($job.start_time,"HH:mm",$null)
          $startFull = $now.Date.Add($startDt.TimeOfDay)
          if ($startFull -gt $now) { $startFull = $startFull.AddDays(-1) }
          $elapsed = [int]($now - $startFull).TotalMinutes
          $limite  = $expected * 2
          if ($elapsed -gt $limite) {
              Write-Host ("[POSSIBLE_HANG] {0,-22} - Transcurrido: {1} min | Límite: {2} min" `
                  -f $job.job_name, $elapsed, $limite)
              $alertas++
          }
      }
  }

  Write-Host ("─" * 62)
  Write-Host ("Jobs procesados: {0} | Alertas generadas: {1}" -f $jobs.Count, $alertas)
  exit $(if ($alertas -gt 0) { 1 } else { 0 })

─────────────────────────────────────────────────────────────────────
  ERRORES COMUNES A OBSERVAR (solo entrevistador):
─────────────────────────────────────────────────────────────────────

  ERROR 1 — No maneja cruce de medianoche
  Síntoma: start_time 22:00, end_time 00:30 da duración negativa
  Impacto: todos los jobs nocturnos lo sufren — error crítico

  ERROR 2 — No valida campos vacíos antes de ParseExact
  Síntoma: ParseExact en end_time="" lanza excepción en runtime
  Impacto: el script falla para cualquier job WAITING o RUNNING

  ERROR 3 — No maneja el exit code según alertas detectadas
  Síntoma: siempre sale con exit 0 aunque haya alertas
  Impacto: Control-M no distingue ejecución limpia de con alertas

  ERROR 4 — Confunde Write-Host con Write-Output
  Síntoma: usa Write-Output y el output va al pipeline en vez
           de a la consola del job de Control-M

  PUNTUACION DEL CODE CHALLENGE:
  +----------------+------------------------------------------------------+
  | 0 - No apto   | No puede escribir el script o no funciona en absoluto|
  | 1 - Básico    | Happy path correcto; falla en casos borde             |
  | 2 - Competente| Campos vacíos + cruce medianoche + exit code correcto |
  | 3 - Destacado | Todo lo anterior + output formateado + explica impacto|
  |               | de cada decisión técnica mientras escribe             |
  +----------------+------------------------------------------------------+

═════════════════════════════════════════════════════════════════════
  SECCIÓN G — SCORECARD RESUMEN
═════════════════════════════════════════════════════════════════════

  TABLA DE PUNTUACIÓN (cada pregunta: 0-3 puntos)

  +-------------+-----+--------------------------------------------------+------+----+
  | Sección     | RE  | Pregunta                                         | Peso | /3 |
  +-------------+-----+--------------------------------------------------+------+----+
  | Orquestac.  |  1  | A-1: Condiciones y estados Control-M             |  15% |    |
  | Orquestac.  |  1  | A-2: Resource pools y distribución de carga      |  12% |    |
  | Orquestac.  |  1  | A-3: Diferencias operativas entre herramientas   |   8% |    |
  | ITIL        |  3  | B-1: Incident response — job crítico a las 03:15 |  15% |    |
  | ITIL        |  3  | B-2: Change Management — migración Control-M     |  10% |    |
  | Diseño      |  4  | C-1: Diseño de malla fan-in/fan-out en vivo      |  15% |    |
  | Diseño      |  4  | C-2: KPIs y observabilidad de la plataforma      |   8% |    |
  | Scripting   |  2  | D-1: Shell script robusto para Control-M         |   7% |    |
  | Diferenciador| D  | E-1: Airflow vs Control-M — coexistencia         |   5% |    |
  | Código      |  2  | F:   Code Challenge PowerShell                   |   5% |    |
  +-------------+-----+--------------------------------------------------+------+----+
  | TOTAL       |     |                                                  | 100% |    |
  +-------------+-----+--------------------------------------------------+------+----+

  FORMULA:
  Score_final (%) = Suma(puntaje_i x peso_i) / 3 x 100

  UMBRALES DE APROBACIÓN — NIVEL SEMI SENIOR:
  +------------------------------------------+---------+-------------------+
  | Resultado                                 | Score   | Condición         |
  +------------------------------------------+---------+-------------------+
  | RECOMENDADO                               | >= 70%  | A-1>=2 Y B-1>=2   |
  | RECOMENDADO PRIORITARIO (diferenciador)  | >= 70%  | + C-1=3 o A-3>=2  |
  | A EVALUAR CON JEFATURA                   | 55-69%  | Sin red flag crit. |
  | NO APTO                                   | < 55%   | O red flag crítico |
  +------------------------------------------+---------+-------------------+

  RED FLAGS QUE RECHAZAN AUTOMÁTICAMENTE:
  X Primera acción ante job fallido: reiniciar sin diagnóstico previo.
  X No distingue entre FAILED y ABEND en Control-M.
  X No tiene plan de rollback concreto para un cambio de plataforma.
  X No puede diseñar dependencias básicas (fan-in/fan-out).
  X No escribe scripts con exit codes explícitos para un scheduler.

─────────────────────────────────────────────────────────────────────
  SEÑALES DE CALIDAD (observar durante toda la entrevista):
─────────────────────────────────────────────────────────────────────

  SEÑAL POSITIVA:
  [ ] Hace preguntas de contexto antes de responder situacionales
  [ ] Admite no saber algo y lo razona en vez de inventar
  [ ] Menciona incidentes reales con detalles técnicos concretos
  [ ] Ajusta su diseño cuando se agrega una restricción nueva
  [ ] Distingue espontáneamente entre jobs idempotentes y no idempotentes

  SEÑAL NEGATIVA:
  [ ] Respuestas que suenan a artículo de blog sin experiencia propia
  [ ] No tiene ningún ejemplo concreto de un fallo o incidente real
  [ ] Llama "cron" a todo tipo de scheduler — vocabulario impreciso
  [ ] No considera el impacto en datos antes de proponer acciones

  OPERACIÓN REGIONAL (Cencosud es multi-país):
  [ ] Demuestra experiencia con múltiples zonas horarias y cierres
      por país — diferenciador relevante para el rol
  [ ] Entiende el modelo de días hábiles por país en los schedulers

  NOTAS DEL ENTREVISTADOR:
  _____________________________________________________________________
  _____________________________________________________________________
  _____________________________________________________________________
  _____________________________________________________________________

  DECISIÓN FINAL:  [ ] Recomendado  [ ] A evaluar  [ ] No apto

  Entrevistador: ________________________________  Fecha: ______________

─────────────────────────────────────────────────────────────────────
  Challenge generado: 2026-05-22 — TechChallenge Architect | ACL Chile
  Diseño anti-IA: follow-ups verbales sorpresa · diseño de mallas en
  vivo con interrupciones y restricciones nuevas · vocabulario operativo
  de nicho (Odate, IN/OUT Conditions, ABEND, resource pool, max wait,
  fan-in/fan-out) que expone respuestas genéricas · incident response
  a las 3am sin runbook · código PowerShell sin IDE ni IA · preguntas
  de experiencia real que exigen narrativa concreta, no definiciones
========================================================================="""


def set_margins(doc, top_cm, bottom_cm, left_cm, right_cm):
    section = doc.sections[0]
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)
    section.left_margin = Cm(left_cm)
    section.right_margin = Cm(right_cm)


def add_plain_paragraph(doc, text, font_name="Courier New", font_size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(font_size * 1.2)

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = False
    run.italic = False
    run.underline = False

    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rPr.insert(0, rFonts)

    return para


def main():
    output_dir = os.path.dirname(OUTPUT_PATH)
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    set_margins(doc, top_cm=1.5, bottom_cm=1.5, left_cm=1.5, right_cm=1.5)

    for line in CONTENT.split("\n"):
        add_plain_paragraph(doc, line)

    doc.save(OUTPUT_PATH)
    print(f"Documento guardado en:\n  {OUTPUT_PATH}")
    print(f"Tamaño: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    main()
